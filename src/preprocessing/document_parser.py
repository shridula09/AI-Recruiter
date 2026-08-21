from pathlib import Path
from pypdf import PdfReader
from docx import Document

def read_text_file(path):
    return Path(path).read_text(encoding="utf-8", errors="ignore")

def parse_document(path_or_uploaded_file):
    name = getattr(path_or_uploaded_file, "name", str(path_or_uploaded_file)).lower()

    if hasattr(path_or_uploaded_file, "read") and not isinstance(path_or_uploaded_file, (str, Path)):
        data = path_or_uploaded_file.read()
        if name.endswith(".txt"):
            return data.decode("utf-8", errors="ignore")
        if name.endswith(".pdf"):
            import io
            reader = PdfReader(io.BytesIO(data))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        if name.endswith(".docx"):
            import io
            doc = Document(io.BytesIO(data))
            return "\n".join(p.text for p in doc.paragraphs)
        raise ValueError("Supported uploads: .txt, .pdf, .docx")

    path = Path(path_or_uploaded_file)
    if path.suffix.lower() == ".txt":
        return read_text_file(path)
    if path.suffix.lower() == ".pdf":
        reader = PdfReader(str(path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    if path.suffix.lower() == ".docx":
        doc = Document(str(path))
        return "\n".join(p.text for p in doc.paragraphs)
    raise ValueError("Supported files: .txt, .pdf, .docx")
